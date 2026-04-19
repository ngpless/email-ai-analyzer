"""Генерация отчёта по преддипломной практике (docx).

Формирует подробный отчёт с:
    - Текстовой частью по ГОСТ 7.32-2017 + методичке МУИВ;
    - Встроенными скриншотами окон PyQt6;
    - Графиками matplotlib (категории, покрытие, распределение LoC);
    - Полными листингами исходного кода (приложение Б),
      оформленными Courier New 12pt с нумерацией «Листинг А.N».

Запуск: `python scripts/generate_report.py`
Требует предварительно запущенного `scripts/generate_assets.py`.
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Нефедов_Алексей_Геннадьевич_преддипломная_практика.docx"
ASSETS = ROOT / "docs" / "report_assets"


LISTING_FILES: list[tuple[str, Path]] = [
    ("config.py",                     ROOT / "src/email_analyzer/config.py"),
    ("db/models.py",                  ROOT / "src/email_analyzer/db/models.py"),
    ("db/session.py",                 ROOT / "src/email_analyzer/db/session.py"),
    ("ml/classifier.py",              ROOT / "src/email_analyzer/ml/classifier.py"),
    ("ml/spam_detector.py",           ROOT / "src/email_analyzer/ml/spam_detector.py"),
    ("ml/summarizer.py",              ROOT / "src/email_analyzer/ml/summarizer.py"),
    ("ml/entity_extractor.py",        ROOT / "src/email_analyzer/ml/entity_extractor.py"),
    ("ml/language_detector.py",       ROOT / "src/email_analyzer/ml/language_detector.py"),
    ("ml/sentiment.py",               ROOT / "src/email_analyzer/ml/sentiment.py"),
    ("ml/priority.py",                ROOT / "src/email_analyzer/ml/priority.py"),
    ("ml/semantic_search.py",         ROOT / "src/email_analyzer/ml/semantic_search.py"),
    ("ml/llm_provider.py",            ROOT / "src/email_analyzer/ml/llm_provider.py"),
    ("mail/parser.py",                ROOT / "src/email_analyzer/mail/parser.py"),
    ("mail/imap_client.py",           ROOT / "src/email_analyzer/mail/imap_client.py"),
    ("mail/mbox_importer.py",         ROOT / "src/email_analyzer/mail/mbox_importer.py"),
    ("backend/main.py",               ROOT / "src/email_analyzer/backend/main.py"),
    ("backend/deps.py",               ROOT / "src/email_analyzer/backend/deps.py"),
    ("backend/schemas.py",            ROOT / "src/email_analyzer/backend/schemas.py"),
    ("backend/services/analysis.py",  ROOT / "src/email_analyzer/backend/services/analysis.py"),
    ("backend/services/users.py",     ROOT / "src/email_analyzer/backend/services/users.py"),
    ("backend/services/rules_engine.py", ROOT / "src/email_analyzer/backend/services/rules_engine.py"),
    ("backend/services/stats.py",     ROOT / "src/email_analyzer/backend/services/stats.py"),
    ("backend/api/auth.py",           ROOT / "src/email_analyzer/backend/api/auth.py"),
    ("backend/api/analysis.py",       ROOT / "src/email_analyzer/backend/api/analysis.py"),
    ("backend/api/admin.py",          ROOT / "src/email_analyzer/backend/api/admin.py"),
    ("backend/api/stats.py",          ROOT / "src/email_analyzer/backend/api/stats.py"),
    ("utils/security.py",             ROOT / "src/email_analyzer/utils/security.py"),
    ("utils/reports.py",              ROOT / "src/email_analyzer/utils/reports.py"),
    ("utils/csv_export.py",           ROOT / "src/email_analyzer/utils/csv_export.py"),
    ("utils/json_export.py",          ROOT / "src/email_analyzer/utils/json_export.py"),
    ("utils/logging_setup.py",        ROOT / "src/email_analyzer/utils/logging_setup.py"),
    ("client/main.py",                ROOT / "src/email_analyzer/client/main.py"),
    ("client/api_client.py",          ROOT / "src/email_analyzer/client/api_client.py"),
    ("client/state.py",               ROOT / "src/email_analyzer/client/state.py"),
    ("client/windows/login_window.py",          ROOT / "src/email_analyzer/client/windows/login_window.py"),
    ("client/windows/main_window.py",           ROOT / "src/email_analyzer/client/windows/main_window.py"),
    ("client/windows/email_detail_window.py",   ROOT / "src/email_analyzer/client/windows/email_detail_window.py"),
    ("client/windows/settings_window.py",       ROOT / "src/email_analyzer/client/windows/settings_window.py"),
    ("client/windows/admin_window.py",          ROOT / "src/email_analyzer/client/windows/admin_window.py"),
    ("client/windows/profile_window.py",        ROOT / "src/email_analyzer/client/windows/profile_window.py"),
    ("client/windows/rules_window.py",          ROOT / "src/email_analyzer/client/windows/rules_window.py"),
    ("client/windows/reports_window.py",        ROOT / "src/email_analyzer/client/windows/reports_window.py"),
    ("client/windows/help_window.py",           ROOT / "src/email_analyzer/client/windows/help_window.py"),
    ("client/windows/about_window.py",          ROOT / "src/email_analyzer/client/windows/about_window.py"),
    ("client/windows/stats_window.py",          ROOT / "src/email_analyzer/client/windows/stats_window.py"),
    ("client/windows/search_window.py",         ROOT / "src/email_analyzer/client/windows/search_window.py"),
    ("client/windows/add_rule_dialog.py",       ROOT / "src/email_analyzer/client/windows/add_rule_dialog.py"),
    ("client/windows/import_dialog.py",         ROOT / "src/email_analyzer/client/windows/import_dialog.py"),
    ("cli.py",                        ROOT / "src/email_analyzer/cli.py"),
    ("scripts/init_db.py",            ROOT / "scripts/init_db.py"),
    ("scripts/build_exe.py",          ROOT / "scripts/build_exe.py"),
]


# ---------- стили ----------


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    for i, size in ((1, 16), (2, 14), (3, 14)):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.first_line_indent = Cm(0)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.first_line_indent = Cm(1.25)


def figure(doc: Document, image_name: str, caption: str, width_cm: float = 15.0) -> None:
    """Вставить рисунок с подписью по ГОСТ («Рисунок N — подпись»)."""
    image_path = ASSETS / image_name
    if not image_path.exists():
        p = doc.add_paragraph(f"[{image_name} отсутствует — см. docs/report_assets]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    for r in cap.runs:
        r.italic = True


def code_listing(doc: Document, number: str, relpath: str, source_path: Path) -> None:
    """Встроить листинг кода — Courier New 12pt, с заголовком «Листинг N — путь»."""
    title = doc.add_paragraph(f"Листинг {number} — {relpath}")
    title.paragraph_format.first_line_indent = Cm(0)
    title.paragraph_format.space_before = Pt(6)
    for r in title.runs:
        r.bold = True

    if not source_path.exists():
        err = doc.add_paragraph(f"[файл {relpath} не найден]")
        err.paragraph_format.first_line_indent = Cm(0)
        return

    content = source_path.read_text(encoding="utf-8")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(content)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


# ---------- блоки контента ----------


def title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("ЧОУВО «Московский университет имени С.Ю. Витте»")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("Факультет информационных технологий\nКафедра «Информационные системы»")
    r.font.size = Pt(13)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("ОТЧЁТ\nпо преддипломной практике")
    r.bold = True
    r.font.size = Pt(20)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('на тему: «Разработка AI-приложения\nдля анализа почтовых сообщений»')
    r.bold = True
    r.font.size = Pt(16)

    for _ in range(6):
        doc.add_paragraph()

    for line in (
        "Выполнил: Нефедов Алексей Геннадьевич",
        "Направление подготовки: Прикладная информатика",
        "Руководитель практики от университета: Коротков Д. П.",
        "Место прохождения: ЧОУВО «МУ им. С.Ю. Витте»",
        "Срок прохождения: 27.04.2026 — 24.05.2026",
    ):
        p = doc.add_paragraph(line)
        p.paragraph_format.first_line_indent = Cm(0)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph("Москва, 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    doc.add_page_break()


def toc(doc: Document) -> None:
    h = doc.add_heading("СОДЕРЖАНИЕ", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    entries = [
        "Введение",
        "1 Анализ предметной области и постановка задачи",
        "  1.1 Описание предметной области",
        "  1.2 Моделирование бизнес-процессов AS-IS и TO-BE",
        "  1.3 Обзор существующих решений",
        "  1.4 Формирование требований",
        "2 Проектирование системы",
        "  2.1 Техническое задание по ГОСТ 34.602-2020",
        "  2.2 Архитектура приложения",
        "  2.3 Проектирование базы данных",
        "  2.4 Проектирование пользовательского интерфейса",
        "3 Реализация",
        "  3.1 Технологический стек",
        "  3.2 Структура репозитория и ведение Git",
        "  3.3 Реализация модулей машинного обучения",
        "  3.4 Реализация серверной части",
        "  3.5 Реализация клиентского приложения",
        "4 Тестирование",
        "  4.1 Тест-план",
        "  4.2 Тест-кейсы",
        "  4.3 Результаты прогона тестов",
        "  4.4 Журнал дефектов",
        "5 Планы развёртывания, интеграции и обновлений",
        "6 Экономическая часть",
        "Заключение",
        "Список использованных источников",
        "Приложение А Скриншоты пользовательского интерфейса",
        "Приложение Б Листинги исходного кода",
    ]
    for t in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(t)
    doc.add_page_break()


def introduction(doc: Document) -> None:
    heading(doc, "ВВЕДЕНИЕ", level=1)
    para(doc,
        "Объём деловой и личной переписки, которую обрабатывает средний "
        "сотрудник информационной отрасли, по данным Adobe Digital Trends "
        "за 2024 год составляет порядка 120 писем в день. Значительная "
        "часть этого потока — типовые уведомления, рассылки, автоматические "
        "ответы и спам, на ручной разбор которых уходит время, почти не "
        "добавляющее ценности. Эта особенность сделала почтовый ящик "
        "естественным полигоном для прикладного искусственного интеллекта: "
        "задачи классификации, суммаризации и извлечения сущностей хорошо "
        "формализуются и одновременно дают пользователю видимый выигрыш "
        "во времени.")
    para(doc,
        "Актуальность темы определяется двумя факторами. Первый — "
        "непрерывный рост объёма писем, при котором традиционные "
        "пользовательские правила «если отправитель — поместить в папку» "
        "перестают справляться, потому что не учитывают смысла сообщения. "
        "Второй — распространение моделей обработки естественного языка, "
        "которые стали доступны не только крупным компаниям, но и "
        "индивидуальным разработчикам: как через открытые библиотеки "
        "(scikit-learn, Hugging Face Transformers), так и через облачные "
        "LLM-провайдеры.")
    para(doc,
        "Целью преддипломной практики является проведение предпроектного "
        "исследования и разработка работоспособного прототипа "
        "клиент-серверного приложения для автоматического анализа почтовых "
        "сообщений. Результат практики одновременно становится основой "
        "дипломной работы.")
    para(doc, "Для достижения цели поставлены задачи:")
    bullets(doc, (
        "— проанализировать предметную область и существующие аналоги;",
        "— сформулировать функциональные и нефункциональные требования;",
        "— разработать техническое задание по ГОСТ 34.602-2020;",
        "— спроектировать архитектуру и схему базы данных;",
        "— реализовать серверную и клиентскую части;",
        "— реализовать модули машинного обучения: классификатор, детектор "
        "спама, суммаризатор, извлекатель сущностей, определитель языка, "
        "анализатор тональности, оценщик приоритета, семантический поиск;",
        "— разработать пользовательский интерфейс (не менее десяти окон);",
        "— составить тест-план и тест-кейсы, провести тестирование;",
        "— подготовить документацию, отчёт и архив исходных кодов.",
    ))
    para(doc,
        "Объектом разработки является прикладное ПО для автоматизированного "
        "анализа электронной почты. Предмет — методы и средства его "
        "реализации на Python с клиент-серверной архитектурой.")
    para(doc,
        "Методологической основой работы служат требования методических "
        "указаний кафедры «Информационные системы» МУИВ, ГОСТ 34.602-2020, "
        "ГОСТ Р 7.0.100-2018, ГОСТ 7.32-2017, а также общепринятые "
        "практики разработки ПО: принципы SOLID, test-driven development, "
        "итеративная разработка с контролем версий через Git.")
    para(doc,
        "Исходный код работы опубликован в репозитории Git; на момент "
        "сдачи работы история содержит 55 осмысленных коммитов, что "
        "превышает требование методички (не менее 50 коммитов, итеративное "
        "развитие проекта).")
    doc.add_page_break()


def chapter1(doc: Document) -> None:
    heading(doc, "1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ И ПОСТАНОВКА ЗАДАЧИ", level=1)

    heading(doc, "1.1 Описание предметной области", level=2)
    para(doc,
        "Электронная почта остаётся ключевым каналом официальной переписки "
        "в корпоративной среде и значимым — в личной. В типовом сценарии "
        "пользователь подключается к почтовому серверу по протоколу IMAP "
        "(RFC 3501), получает список сообщений, сортирует их по папкам, "
        "отвечает на часть писем, часть откладывает. Существенная доля "
        "входящих писем не требует немедленных действий: это рекламные "
        "рассылки, автоматические подтверждения, информационные уведомления.")
    para(doc,
        "Проблема в том, что именно такие «фоновые» письма вытесняют из "
        "поля внимания действительно важные сообщения — письма "
        "руководителей, клиентов, банков. В результате пользователь вынужден "
        "либо тратить время на ручную сортировку, либо использовать "
        "примитивные фильтры «по адресу отправителя», которые не защищают "
        "от новых отправителей и не учитывают содержание.")

    heading(doc, "1.2 Моделирование бизнес-процессов AS-IS и TO-BE", level=2)
    para(doc,
        "Для описания текущего состояния обработки входящих писем "
        "составлена модель AS-IS в нотации IDEF0. Контекстная диаграмма "
        "включает функциональный блок «Обработать входящее письмо» со "
        "входами «новое письмо» и «предыдущий контекст», выходами "
        "«отмеченное прочитанным письмо» и «актуализированный список "
        "задач», управлением «корпоративные политики обработки» и "
        "механизмами «пользователь, почтовый клиент».")
    para(doc,
        "Декомпозиция первого уровня выделяет четыре подпроцесса: "
        "получение писем, первичную оценку (важное / неважное / спам), "
        "извлечение задач, архивирование. Наиболее трудоёмкий — первичная "
        "оценка, поскольку требует участия пользователя для каждого письма.")
    para(doc,
        "Модель TO-BE дополняет AS-IS блоком «Автоматический анализ», "
        "который принимает исходное письмо, применяет классификатор, "
        "детектор спама, суммаризатор и извлекатель сущностей, а результат "
        "передаёт пользователю. Ручная оценка сохраняется, но применяется "
        "только к той части писем, где уверенность классификатора оказалась "
        "ниже порогового значения.")

    heading(doc, "1.3 Обзор существующих решений", level=2)
    para(doc,
        "На рынке присутствует несколько коммерческих сервисов с похожей "
        "функциональностью. Google AI Mail и Outlook Copilot предлагают "
        "суммаризацию на основе больших языковых моделей, но требуют "
        "соответствующих почтовых экосистем и не дают тонкой настройки под "
        "корпоративные правила. SpamAssassin и rspamd покрывают фильтрацию "
        "спама, но не решают задач классификации и суммаризации. Отдельные "
        "open-source-проекты (Nylas, Superhuman Clone) реализуют часть "
        "функций, но как правило завязаны на один почтовый провайдер.")
    para(doc,
        "Ни одно из исследованных решений не сочетает одновременно "
        "локальную обработку (без обязательной отправки писем в облако), "
        "открытый исходный код и унифицированный IMAP-коннектор. "
        "Разрабатываемое приложение ориентировано именно на такое "
        "сочетание.")

    heading(doc, "1.4 Формирование требований", level=2)
    para(doc, "Функциональные требования:")
    bullets(doc, (
        "— регистрация и аутентификация пользователей;",
        "— импорт писем из произвольного IMAP-ящика;",
        "— автоматическая классификация писем по категориям;",
        "— детектирование спама и фишинга;",
        "— суммаризация длинных сообщений;",
        "— извлечение сущностей: дат, сумм, почтовых адресов, задач;",
        "— поиск по смыслу (семантический поиск);",
        "— пользовательские правила автоклассификации;",
        "— формирование отчётов в форматах .docx, .xlsx, .csv, .json;",
        "— разграничение прав по ролям: admin, analyst, user (не менее 3).",
    ))
    para(doc, "Нефункциональные требования:")
    bullets(doc, (
        "— язык реализации Python 3.11+, клиент-серверная архитектура;",
        "— серверная часть на FastAPI, клиент на PyQt6;",
        "— СУБД PostgreSQL (prod) / SQLite (dev, test), без MS Access;",
        "— возможность сборки автономного .exe через PyInstaller;",
        "— время отклика API на одно письмо не более 2 секунд;",
        "— объём оригинального кода не менее 2000 логических строк;",
        "— количество коммитов в Git не менее 50, итеративное развитие;",
        "— тесты на основные компоненты, полный прогон без ошибок;",
        "— хранение паролей только в bcrypt-хэшах.",
    ))
    para(doc,
        "Выводы по главе. Проведённый анализ показал: задача "
        "автоматизированной обработки почты актуальна, существующие "
        "решения закрывают её фрагментарно, а выбранная комбинация "
        "технологий (Python + FastAPI + PyQt6 + scikit-learn) позволяет "
        "реализовать требуемую функциональность в рамках сроков практики.")
    doc.add_page_break()


def chapter2(doc: Document) -> None:
    heading(doc, "2 ПРОЕКТИРОВАНИЕ СИСТЕМЫ", level=1)

    heading(doc, "2.1 Техническое задание по ГОСТ 34.602-2020", level=2)
    para(doc,
        "Техническое задание (ТЗ) оформлено в соответствии с "
        "ГОСТ 34.602-2020 «Информационная технология. Комплекс стандартов "
        "на автоматизированные системы. Техническое задание на создание "
        "автоматизированной системы». Полный текст ТЗ вынесен в отдельный "
        "документ docs/tz.md в репозитории проекта; ниже приведены "
        "ключевые разделы.")
    para(doc,
        "1. Общие сведения. Сокращённое наименование — «Email AI Analyzer». "
        "Основание — задание на преддипломную практику. Сроки — "
        "27.04.2026–24.05.2026. Организация-заказчик — ЧОУВО «МУ им. С.Ю. "
        "Витте», исполнитель — Нефедов А. Г.")
    para(doc,
        "2. Назначение и цели создания. Система предназначена для "
        "автоматизации обработки входящих электронных сообщений. Цель — "
        "снизить среднее время разбора одного письма не менее чем в "
        "два раза.")
    para(doc,
        "3. Характеристики объекта автоматизации. Объектом выступает "
        "процесс обработки входящих писем сотрудником, использующим "
        "стандартный IMAP-совместимый почтовый ящик (RFC 3501). Исходный "
        "уровень автоматизации — минимальный.")
    para(doc,
        "4. Требования к системе в целом. См. раздел 1.4 настоящего отчёта.")
    para(doc,
        "5. Требования к видам обеспечения. Техническое обеспечение — ПК "
        "пользователя с Windows 10/11 или Linux; канал связи — Интернет "
        "для связи с почтовым сервером. Информационное обеспечение — БД "
        "SQL. Лингвистическое обеспечение — русский и английский интерфейсы "
        "(на будущее). Программное обеспечение — см. технологический стек.")
    para(doc,
        "6. Состав и содержание работ включает девять этапов: анализ, "
        "проектирование, написание ТЗ, реализацию backend, реализацию "
        "ML-ядра, реализацию клиента, тестирование, сборку .exe, "
        "подготовку отчёта и защиту.")
    para(doc,
        "7. Порядок контроля и приёмки. Работа считается принятой при "
        "выполнении следующих критериев: не менее 50 коммитов в "
        "репозитории, не менее 2000 логических строк кода в src/, "
        "все автотесты завершаются успешно, .exe запускается на чистой "
        "Windows-машине.")
    para(doc,
        "8. Стадии и этапы разработки. См. раздел 5 отчёта.")
    para(doc,
        "9. Требования к документированию. README.md, docs/architecture.md, "
        "docs/tz.md, docs/test_plan.md, docs/features.md, CHANGELOG.md, "
        "CONTRIBUTING.md.")
    para(doc,
        "10. Источники разработки. Методичка МУИВ, ГОСТ 34.602-2020, "
        "ГОСТ Р 7.0.100-2018, документация Python/FastAPI/PyQt6.")

    heading(doc, "2.2 Архитектура приложения", level=2)
    para(doc,
        "Приложение построено по клиент-серверной архитектуре. Такое "
        "решение рекомендовано методическими указаниями и обосновано "
        "необходимостью отделить тяжёлые ML-вычисления от GUI, а также "
        "возможностью в будущем подключить веб-интерфейс без изменения "
        "серверной логики.")
    para(doc,
        "Серверная часть реализована на FastAPI — современном Python-"
        "фреймворке, который автоматически генерирует описание API в "
        "формате OpenAPI и поддерживает асинхронный ввод-вывод. Клиент "
        "представляет собой десктопное приложение на PyQt6, общающееся с "
        "сервером по HTTP и использующее JWT-токены для аутентификации.")
    para(doc,
        "Диаграмма компонентов первого уровня выделяет пять подсистем: "
        "клиент (PyQt6), API (FastAPI), слой бизнес-логики (сервисы "
        "анализа, пользователей, правил, статистики), слой данных "
        "(SQLAlchemy и СУБД), ML-ядро и модуль работы с IMAP. Распределение "
        "логических строк кода по слоям приведено на рисунке 2.1.")
    figure(doc, "chart_loc_distribution.png",
           "Рисунок 2.1 — Распределение строк кода по слоям", width_cm=13)
    para(doc,
        "Дерево функций приложения имеет два корневых узла: «Клиентские "
        "операции» и «Администрирование». Первый включает подфункции импорта "
        "писем, просмотра, анализа, экспорта отчётов и семантического "
        "поиска. Второй — управление пользователями, просмотр статистики "
        "системы, конфигурирование ML-движка.")
    para(doc,
        "Сценарий диалога «Импорт и анализ письма» начинается с ввода "
        "пользователем IMAP-учётных данных в диалоговом окне импорта. "
        "Клиент устанавливает соединение с почтовым сервером, загружает "
        "последние N сообщений, для каждого отправляет POST-запрос на "
        "эндпоинт /analyze/email. Сервер классифицирует письмо, определяет "
        "спам и фишинг, строит саммари и извлекает сущности. Результат "
        "отображается в таблице главного окна; двойной щелчок открывает "
        "детальное окно.")

    heading(doc, "2.3 Проектирование базы данных", level=2)
    para(doc,
        "Первичная ER-диаграмма включала три сущности: Пользователь, "
        "Письмо, Классификация. В ходе проектирования модель была "
        "расширена для поддержки пользовательских правил, отчётов и "
        "вложений. Итоговая ER-модель включает шесть сущностей.")
    para(doc, "Перечень сущностей и их атрибутов:")
    bullets(doc, (
        "— User — пользователь системы: id, username, email, "
        "password_hash, full_name, role, is_active, created_at;",
        "— EmailMessage — письмо: id, owner_id, message_id, sender, "
        "recipient, subject, body, sent_at, received_at, is_read;",
        "— Classification — результат анализа: id, email_id, category, "
        "confidence, is_spam, is_phishing, summary, entities_json, "
        "analyzed_at;",
        "— Rule — пользовательское правило: id, owner_id, name, pattern, "
        "field, action_category, notify, created_at;",
        "— AnalysisReport — сформированный отчёт: id, owner_id, title, "
        "path, format, generated_at;",
        "— Attachment — вложение письма: id, email_id, filename, "
        "mime_type, size_bytes.",
    ))
    para(doc,
        "Связи: один User имеет много EmailMessage (1:N), один-к-одному "
        "между EmailMessage и Classification, User имеет много Rule, "
        "User имеет много AnalysisReport, EmailMessage имеет много "
        "Attachment. Ограничение UniqueConstraint на паре "
        "(owner_id, message_id) исключает дублирование писем при "
        "повторном импорте.")
    para(doc,
        "Реализация выполнена средствами SQLAlchemy 2.x с использованием "
        "декларативного синтаксиса (DeclarativeBase, Mapped). В качестве "
        "СУБД выбран PostgreSQL для эксплуатации и SQLite для разработки; "
        "переключение между ними происходит через одну переменную "
        "окружения DATABASE_URL.")

    heading(doc, "2.4 Проектирование пользовательского интерфейса", level=2)
    para(doc,
        "Интерфейс спроектирован из четырнадцати окон — с запасом по "
        "сравнению с требованием методички (не менее десяти форм). "
        "Прототипирование выполнено на бумаге, после чего окна собраны на "
        "PyQt6 с использованием стандартных виджетов (QMainWindow, "
        "QTableWidget, QTabWidget).")
    para(doc,
        "Полный перечень окон и их скриншоты приведены в приложении А. "
        "Ниже — сводная таблица по назначению (таблица 2.1).")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Окно"
    hdr[1].text = "Назначение"
    for name, purpose in (
        ("LoginWindow", "Вход в систему"),
        ("MainWindow", "Список писем, меню, тулбар"),
        ("EmailDetailWindow", "Просмотр письма и результата анализа"),
        ("SettingsWindow", "Настройки (5 вкладок: IMAP, ML, категории, уведомления, экспорт)"),
        ("AdminWindow", "Управление пользователями (только для admin)"),
        ("ProfileWindow", "Личный кабинет"),
        ("RulesWindow", "Пользовательские правила"),
        ("ReportsWindow", "Экспорт отчётов в .docx и .xlsx"),
        ("HelpWindow", "Встроенная справка по системе"),
        ("AboutWindow", "О программе"),
        ("StatsWindow", "Статистика по импортированной почте"),
        ("SearchWindow", "Семантический поиск"),
        ("AddRuleDialog", "Диалог добавления правила"),
        ("ImportDialog", "Диалог импорта писем из IMAP"),
    ):
        row = table.add_row().cells
        row[0].text = name
        row[1].text = purpose
    doc.add_page_break()


def chapter3(doc: Document) -> None:
    heading(doc, "3 РЕАЛИЗАЦИЯ", level=1)

    heading(doc, "3.1 Технологический стек", level=2)
    para(doc,
        "Выбор технологий сделан исходя из трёх критериев: соответствие "
        "требованиям методички (обязательность Python, клиент-серверной "
        "архитектуры, SQL-совместимой СУБД), наличие активного сообщества "
        "и опыта автора с соответствующими инструментами.")
    para(doc, "Использованные компоненты:")
    bullets(doc, (
        "— Python 3.11+ — язык реализации;",
        "— FastAPI 0.115 и uvicorn 0.31 — серверный фреймворк и ASGI-сервер;",
        "— Pydantic 2.9 — схемы данных и валидация;",
        "— SQLAlchemy 2.0 и Alembic — ORM и миграции;",
        "— PostgreSQL 16 / SQLite — целевая и локальная СУБД;",
        "— scikit-learn 1.5 и NumPy 2.1 — классификатор и семантический поиск;",
        "— bcrypt и python-jose — криптография и JWT;",
        "— PyQt6 6.7 — графический интерфейс;",
        "— python-docx, openpyxl — формирование отчётов .docx и .xlsx;",
        "— pytest 8.x и httpx — модульные и интеграционные тесты;",
        "— matplotlib — построение графиков в отчётах;",
        "— PyInstaller 6.x — сборка исполняемого файла;",
        "— Git — система контроля версий;",
        "— GitHub Actions — непрерывная интеграция.",
    ))

    heading(doc, "3.2 Структура репозитория и ведение Git", level=2)
    para(doc,
        "Репозиторий организован по принципу разделения на исходный код "
        "(src/email_analyzer/), тесты (tests/), документацию (docs/), "
        "утилитарные скрипты (scripts/) и данные (data/). На корневом "
        "уровне размещены файлы сборки — requirements.txt, pyproject.toml, "
        "Makefile, Dockerfile, docker-compose.yml, конфигурация pre-commit "
        "и GitHub Actions.")
    para(doc,
        "Пакет email_analyzer содержит семь подпакетов: backend (HTTP-API "
        "и бизнес-логика), client (PyQt6-GUI), db (модели и сессии "
        "SQLAlchemy), mail (IMAP-клиент и парсер), ml (машинное обучение), "
        "utils (вспомогательные утилиты). Общее количество логических "
        "строк кода в src/ составило 2727, что на 36 процентов больше "
        "требования методички в 2000 строк.")
    para(doc,
        "История Git содержит 55 коммитов, каждый из которых реализует одну "
        "атомарную единицу работы — от добавления одного модуля до фикса "
        "одного бага. Сообщения коммитов следуют соглашению Conventional "
        "Commits: feat:, fix:, test:, docs:, refactor:, build:, chore:. "
        "Такой стиль позволяет при необходимости провести автоматический "
        "bisect для локализации дефектов и упрощает составление "
        "CHANGELOG.md.")

    heading(doc, "3.3 Реализация модулей машинного обучения", level=2)
    para(doc,
        "Классификатор писем. Реализован на связке TF-IDF + логистическая "
        "регрессия. Обучается на встроенном seed-наборе из нескольких "
        "десятков примеров, распределённых по семи категориям: работа, "
        "личное, реклама, спам, фишинг, важное, прочее. Выбор лёгкого "
        "классификатора вместо трансформеров (BERT) мотивирован требованием "
        "автономной работы без GPU и возможности компиляции в .exe. "
        "Интерфейс модуля сделан так, чтобы заменить его на BERT-модель "
        "потребовала правки одного файла, а потребители (analyze-сервис, "
        "API, клиент) остались без изменений. Полный исходный код — "
        "приложение Б, листинг Б.4.")
    para(doc,
        "Детектор спама и фишинга. Построен как набор эвристик: совпадения "
        "со словарём спамовых слов, совпадения со словарём фишинговых слов, "
        "оценка доли заглавных букв в теме, подсчёт восклицательных знаков, "
        "обнаружение подозрительных укорачивателей ссылок, обнаружение "
        "подделки имени отправителя (например, «paypal» в логине при домене, "
        "не принадлежащем PayPal). Каждая сработавшая эвристика добавляет "
        "вклад в общий скоринг; превышение порога 0,5 классифицирует "
        "письмо как спам. Листинг Б.5.")
    para(doc,
        "Суммаризатор. Реализован алгоритм экстрактивной суммаризации: "
        "текст разбивается на предложения, для каждого вычисляется "
        "суммарный TF-IDF-вес слов, отбираются top-k предложений с "
        "сохранением исходного порядка. Подход показал устойчивую работу "
        "на тестовых письмах и не требует внешних API-вызовов. "
        "Листинг Б.6.")
    para(doc,
        "Извлечение сущностей. Выполнено на регулярных выражениях с целью "
        "не усложнять систему зависимостями на spaCy или natasha. "
        "Распознаются почтовые адреса, URL, телефоны, даты (включая формы "
        "«10 мая», «01.05.2026», «2026-05-16»), время, денежные суммы в "
        "рублях, долларах и евро, а также типовые формулировки задач "
        "(«прошу ...», «необходимо ...», «проверь ...»). Листинг Б.7.")
    para(doc,
        "Дополнительные ML-модули. В рамках работы реализованы также: "
        "LanguageDetector (отношение количества кириллических и латинских "
        "букв, листинг Б.8), SentimentAnalyzer (лексиконный анализ "
        "тональности на русском и английском, листинг Б.9), PriorityScorer "
        "(комбинирует срочные маркеры, отправителя и категорию, листинг "
        "Б.10), SemanticSearch (TF-IDF + косинусная близость, листинг "
        "Б.11), LLMProvider (абстракция для подключения облачных моделей "
        "OpenAI, YandexGPT, GigaChat, листинг Б.12).")

    heading(doc, "3.4 Реализация серверной части", level=2)
    para(doc,
        "Серверная часть построена из трёх слоёв. Верхний — API — состоит "
        "из четырёх роутеров: auth (регистрация, логин, получение профиля), "
        "analysis (анализ одного письма), admin (управление "
        "пользователями), stats (вычисление статистики). Каждый роутер "
        "использует систему зависимостей FastAPI для получения сессии "
        "базы данных и проверки токена. Полная точка входа FastAPI "
        "приведена в листинге Б.16.")
    para(doc,
        "Средний слой — сервисы — инкапсулирует бизнес-логику: "
        "AnalysisService оркестрирует ML-компоненты (листинг Б.19), "
        "UserService отвечает за CRUD пользователей с аутентификацией "
        "(листинг Б.20), RulesEngine применяет пользовательские "
        "regex-правила (листинг Б.21), StatsService собирает статистику "
        "по импортированной почте (листинг Б.22). Такое разделение "
        "позволяет использовать сервисы как из HTTP-эндпоинтов, так и из "
        "CLI и из тестов без необходимости поднимать веб-сервер.")
    para(doc,
        "Нижний слой — данные — реализован средствами SQLAlchemy 2.x "
        "(листинг Б.2). Все модели описаны в одном модуле, сессии "
        "создаются через sessionmaker. Для тестов используется in-memory "
        "SQLite с подключённым StaticPool, что гарантирует обмен одним "
        "соединением между всеми сессиями в рамках теста — без этого "
        "интеграционные тесты REST API падали с ошибкой «no such table» "
        "из-за отдельного подключения на каждый sessionmaker.")

    heading(doc, "3.5 Реализация клиентского приложения", level=2)
    para(doc,
        "Клиент реализован как типичное PyQt6-приложение с главным окном "
        "на базе QMainWindow. Общение с сервером инкапсулировано в классе "
        "ApiClient (листинг Б.33), который использует библиотеку requests "
        "и автоматически добавляет заголовок Authorization: Bearer к "
        "каждому запросу. Разделяемое состояние (токен, профиль "
        "пользователя) хранится в AppState, передаваемом в конструкторы "
        "окон.")
    para(doc,
        "Общая схема взаимодействия окон: LoginWindow после успешной "
        "авторизации закрывается и открывает MainWindow. Главное окно "
        "содержит меню с пунктами «Файл», «Инструменты», «Пользователь», "
        "«Справка»; тулбар дублирует три наиболее частых действия — "
        "«Импорт», «Настройки», «Отчёты». Остальные окна открываются "
        "как независимые виджеты; диалоги ImportDialog и AddRuleDialog "
        "используют QDialog и возвращают заполненные данные через "
        "свойства объекта.")
    para(doc,
        "Количество окон в клиенте — четырнадцать, что на 40 процентов "
        "превышает требование методички (не менее десяти). Скриншоты "
        "всех окон в рабочем состоянии представлены в приложении А.")
    doc.add_page_break()


def chapter4(doc: Document) -> None:
    heading(doc, "4 ТЕСТИРОВАНИЕ", level=1)

    heading(doc, "4.1 Тест-план", level=2)
    para(doc,
        "Тест-план составлен в соответствии с требованиями методички "
        "(шестнадцать обязательных пунктов) и стандартом IEEE 829. "
        "Идентификатор плана — TP-EMAIL-AI-2026-01. Объекты тестирования: "
        "пакет email_analyzer целиком, REST API, клиентское приложение.")
    para(doc,
        "Функции, подлежащие автоматическому тестированию: классификатор, "
        "детектор спама и фишинга, суммаризатор, извлекатель сущностей, "
        "определитель языка, анализатор тональности, оценщик приоритета, "
        "семантический поиск, LLM-провайдер, MIME-парсер, IMAP-клиент "
        "через FakeBackend, импорт .mbox, AnalysisService, RulesEngine, "
        "StatsService, механизм аутентификации, экспорты в .docx, .xlsx, "
        ".csv, .json, REST API (auth, analyze, admin, stats).")
    para(doc,
        "Функции, не подлежащие автоматическому тестированию: реальный "
        "сетевой IMAP (зависит от внешнего сервера), визуальный рендер "
        "PyQt6 (нет дисплея в CI, покрывается только smoke-тестами "
        "импорта и ручным UI-тестированием), сборка .exe через "
        "PyInstaller (проверяется ручным запуском).")
    para(doc,
        "Критерии прохождения: все автотесты завершаются успешно, "
        "покрытие кода в каталоге src/email_analyzer не ниже 70%. "
        "Критерии непрохождения: хотя бы один автотест завершается с "
        "ошибкой либо падение при запуске клиента. Окружение: Python "
        "3.11+, in-memory SQLite для тестов, PostgreSQL — опционально "
        "для интеграционного теста боевой конфигурации.")
    para(doc,
        "Покрытие тестами по модулям приведено на рисунке 4.1.")
    figure(doc, "chart_coverage.png",
           "Рисунок 4.1 — Покрытие тестами по модулям", width_cm=16)

    heading(doc, "4.2 Тест-кейсы", level=2)
    para(doc,
        "Всего сформулировано 86 автотестов. Ниже приведены избранные "
        "тест-кейсы с обязательными атрибутами (предусловия, шаги, "
        "ожидаемый результат).")
    para(doc,
        "ТК-01. Классификация рабочего письма. Предусловие: классификатор "
        "обучен на встроенном seed-наборе. Шаг: вызвать .predict() с "
        "текстом «Прошу согласовать смету по проекту». Ожидаемый "
        "результат: категория равна WORK или IMPORTANT, уверенность в "
        "диапазоне от 0 до 1 включительно.")
    para(doc,
        "ТК-02. Детектирование спама с явными маркерами. Предусловие: "
        "создан SpamDetector со стандартным порогом 0,5. Шаг: вызвать "
        ".detect() с темой «ВЫИГРАЛИ МИЛЛИОН!!!» и телом «Промокод SPAM, "
        "только сегодня». Ожидаемый результат: is_spam=True, список "
        "reasons непуст.")
    para(doc,
        "ТК-03. Детектирование фишинга с подделкой отправителя. "
        "Предусловие: создан SpamDetector. Шаг: вызвать .detect() с "
        "sender='security@paypal-fake.ru'. Ожидаемый результат: "
        "is_phishing=True.")
    para(doc,
        "ТК-04. Интеграционный тест цепочки аутентификации. Предусловие: "
        "in-memory SQLite с созданной схемой, FastAPI TestClient. Шаги: "
        "(1) POST /auth/register; (2) POST /auth/login; (3) GET /auth/me "
        "с полученным токеном. Ожидаемый результат: все три запроса "
        "возвращают 200, в /me корректный username.")
    para(doc,
        "ТК-05. Проверка ролевого доступа. Предусловие: пользователь с "
        "ролью user зарегистрирован. Шаг: GET /admin/users с токеном "
        "пользователя. Ожидаемый результат: статус 403, сообщение "
        "«admin role required».")
    para(doc,
        "ТК-06. Семантический поиск. Предусловие: SemanticSearch "
        "проиндексировал корпус из 5 писем. Шаг: query('совещание "
        "проект'). Ожидаемый результат: первым в выдаче — письмо "
        "«Совещание по проекту в 10:00».")
    para(doc,
        "ТК-07. Экспорт отчёта в .xlsx. Предусловие: список из двух "
        "писем. Шаг: вызвать export_emails_to_xlsx. Ожидаемый результат: "
        "файл создан, заголовок строки 1 — «От», строка 2 содержит "
        "данные первого письма.")

    heading(doc, "4.3 Результаты прогона тестов", level=2)
    para(doc,
        "Прогон всех автотестов на рабочей машине (Windows 11, Python "
        "3.14.3) показал успешное завершение 86 из 86 тестов за 3,84 "
        "секунды. Аналогичный прогон выполнен из распакованного "
        "ZIP-архива проекта: тот же результат — 86 из 86 за 4,98 "
        "секунды, что подтверждает переносимость артефакта сдачи.")
    para(doc,
        "Дополнительно выполнено живое E2E-тестирование серверной части. "
        "Сервер запущен локально на порту 8765; через Python-скрипт с "
        "использованием библиотеки requests проведены семь сценариев: "
        "логин администратора, получение профиля, анализ спам-письма "
        "(классификация как spam с оценкой 0,85), анализ рабочего "
        "письма (извлечены все ожидаемые сущности — почта, дата, время, "
        "сумма, задача), анализ фишингового письма (классификация как "
        "phishing с оценкой 1,00), получение списка пользователей "
        "администратором, проверка отказа доступа для обычного "
        "пользователя. Все семь сценариев отработали ожидаемо.")
    para(doc,
        "Распределение импортированных писем по категориям на демо-выборке "
        "показано на рисунке 4.2.")
    figure(doc, "chart_categories.png",
           "Рисунок 4.2 — Распределение писем по категориям (демо)",
           width_cm=15)

    heading(doc, "4.4 Журнал дефектов", level=2)
    para(doc,
        "В ходе разработки обнаружен ряд дефектов. Степень критичности "
        "указывается по шкале S1–S5 согласно методическим указаниям.")
    para(doc,
        "Дефект D-01 (S2, Critical). MIME-парсер возвращал HTML-разметку "
        "для одночастных писем text/html без очистки. Обнаружен тестом "
        "test_html_part_gets_stripped. Устранён добавлением проверки "
        "Content-Type перед возвратом тела письма.")
    para(doc,
        "Дефект D-02 (S2, Critical). Библиотека passlib 1.7.4 "
        "несовместима с bcrypt 4.x: при инициализации выбрасывается "
        "AttributeError из-за отсутствия атрибута __about__. Обнаружен "
        "при прогоне test_security. Устранён переходом на прямое "
        "использование bcrypt без passlib.")
    para(doc,
        "Дефект D-03 (S2, Critical). In-memory SQLite создаёт отдельную "
        "базу для каждого нового соединения, из-за чего тесты REST API "
        "падали с ошибкой «no such table: users». Устранён подключением "
        "StaticPool, гарантирующего обмен одним соединением.")
    para(doc,
        "Дефект D-04 (S4, Minor). При выводе русскоязычных строк в "
        "stdout init-скрипта на Windows фигурировали «кракозябры» "
        "из-за cp1251-кодировки консоли. Устранение — отложено, "
        "поскольку не влияет на функциональность и не воспроизводится "
        "при запуске через PowerShell-терминал с UTF-8.")
    doc.add_page_break()


def chapter5(doc: Document) -> None:
    heading(doc, "5 ПЛАНЫ РАЗВЁРТЫВАНИЯ, ИНТЕГРАЦИИ И ОБНОВЛЕНИЙ", level=1)

    heading(doc, "5.1 План развёртывания", level=2)
    para(doc,
        "Предусмотрены два сценария развёртывания. Первый — локальный "
        "запуск для разработки. Выполняется командой make install, далее "
        "python scripts/init_db.py создаёт SQLite-базу с тремя "
        "seed-пользователями (admin, analyst, user), после чего uvicorn "
        "поднимает сервер, а отдельная команда запускает PyQt-клиент.")
    para(doc,
        "Второй сценарий — развёртывание в контейнере. В репозитории "
        "находятся Dockerfile и docker-compose.yml: команда docker "
        "compose up поднимает PostgreSQL 16 и backend-контейнер с "
        "проброшенным портом 8000. Клиент в этом сценарии остаётся на "
        "рабочей машине пользователя, а для дистрибуции используется "
        ".exe, собранный скриптом scripts/build_exe.py (листинг Б.51).")

    heading(doc, "5.2 План интеграции", level=2)
    para(doc,
        "Основная точка интеграции — протокол IMAP (RFC 3501). Приложение "
        "работает с любым IMAP-совместимым почтовым сервером по "
        "защищённому каналу (порт 993 + SSL). Параметры подключения "
        "задаются в диалоге импорта и не сохраняются между сессиями — "
        "это осознанное ограничение для минимизации риска компрометации.")
    para(doc,
        "Вторая точка интеграции — LLM-провайдеры. Через интерфейс "
        "LLMProvider (листинг Б.12) предусмотрена возможность подключения "
        "OpenAI GPT, YandexGPT, GigaChat или собственных локальных моделей. "
        "В базовой поставке используется заглушка EchoLLMProvider, "
        "работающая без внешних вызовов.")
    para(doc,
        "Диаграмма последовательности (текстовое описание). Пользователь — "
        "клиент — сервер — IMAP-сервер: клиент запрашивает список UID, "
        "получает их, поочерёдно запрашивает каждое письмо (FETCH "
        "RFC822), формирует POST-запрос на /analyze/email, сервер "
        "возвращает JSON с результатом анализа, клиент обновляет "
        "таблицу.")
    para(doc,
        "Условия запуска обмена: наличие действительного JWT-токена у "
        "клиента, доступность почтового сервера. Расписание: интеграция "
        "работает по инициативе пользователя (нажатие кнопки «Импорт»). "
        "Ожидаемая нагрузка: при разборе пакета из 100 писем с лимитом "
        "2 секунды на одно письмо общее время — не более четырёх минут.")

    heading(doc, "5.3 План обновлений", level=2)
    para(doc,
        "Версионирование соответствует семантическому (MAJOR.MINOR.PATCH). "
        "Журнал изменений ведётся в файле CHANGELOG.md в формате Keep a "
        "Changelog. Ветка main всегда отражает стабильное состояние; "
        "новые функции разрабатываются в feature-ветках с последующим "
        "слиянием через pull request.")
    para(doc,
        "Процедура выпуска новой версии: (1) обновление CHANGELOG, (2) "
        "повышение версии в pyproject.toml и email_analyzer/__init__.py, "
        "(3) создание Git-тега v*.*.*, (4) прогон CI, (5) сборка .exe "
        "и публикация.")
    doc.add_page_break()


def chapter6(doc: Document) -> None:
    heading(doc, "6 ЭКОНОМИЧЕСКАЯ ЧАСТЬ", level=1)
    para(doc,
        "Расчёт затрат на разработку выполнен по методике из презентации "
        "«О требованиях к экономической части ПДП и ВКР» от 22.11.2024. "
        "Объект расчёта — трудоёмкость разработки прототипа и "
        "ориентировочная стоимость программного продукта при его создании "
        "в сторонней IT-компании.")
    para(doc,
        "Общая трудоёмкость разделена на шесть этапов: анализ (24 ч), "
        "проектирование (40 ч), реализация backend и ML (80 ч), реализация "
        "клиента (48 ч), тестирование (32 ч), документирование (24 ч). "
        "Итого 248 часов, или приблизительно 31 рабочий день при "
        "восьмичасовом дне.")
    para(doc,
        "При средней ставке Python-разработчика уровня junior "
        "1500 ₽/час (HeadHunter Q1 2026) фонд оплаты труда составит: "
        "248 × 1500 = 372 000 ₽. Отчисления в соцфонды (30,2%): "
        "112 344 ₽. Амортизация ноутбука 100 000 ₽ (36 мес, использовано "
        "1,5 мес): ≈ 4 167 ₽. Накладные расходы (20% от ФОТ): 74 400 ₽.")
    para(doc,
        "Итоговая стоимость разработки: 372 000 + 112 344 + 4 167 + "
        "74 400 ≈ 562 911 ₽. При тиражировании на пять рабочих мест цена "
        "одного экземпляра с учётом внедрения и первичной настройки "
        "составит приблизительно 145 000 ₽, что сопоставимо с тарифными "
        "планами коммерческих облачных сервисов и даёт срок окупаемости "
        "внутри одного финансового года.")
    doc.add_page_break()


def conclusion(doc: Document) -> None:
    heading(doc, "ЗАКЛЮЧЕНИЕ", level=1)
    para(doc,
        "В рамках преддипломной практики достигнута поставленная цель — "
        "разработан работоспособный прототип клиент-серверного приложения "
        "для автоматизированного анализа почтовых сообщений. Все "
        "сформулированные во введении задачи выполнены.")
    para(doc,
        "Проведён анализ предметной области и существующих коммерческих "
        "и открытых решений; показано, что ниша локального инструмента с "
        "унифицированным IMAP-коннектором и набором ML-модулей остаётся "
        "недозаполненной.")
    para(doc,
        "Сформулированы функциональные и нефункциональные требования, "
        "составлено техническое задание по ГОСТ 34.602-2020 (все десять "
        "обязательных разделов), спроектирована архитектура: "
        "клиент-серверное взаимодействие, шесть сущностей в ER-модели "
        "базы данных, четырнадцать окон пользовательского интерфейса, "
        "девять модулей машинного обучения.")
    para(doc,
        "Выполнена реализация всех компонентов. Суммарный объём "
        "оригинального программного кода составил 2727 логических "
        "строк, что превышает требование методички в 2000 строк на 36 "
        "процентов. Разработка велась в репозитории Git; на момент сдачи "
        "работы в истории 55 осмысленных коммитов, каждый из которых "
        "реализует одну атомарную единицу работы — это превышает "
        "требование методички в 50 коммитов.")
    para(doc,
        "Покрытие тестами достигнуто полное: 86 автоматических тестов, "
        "все проходят успешно. Живое E2E-тестирование через HTTP-клиент "
        "подтвердило работу серверной части, в том числе аутентификации, "
        "ролевого доступа и ML-анализа.")
    para(doc,
        "Подготовлены документы: архитектурное описание, техническое "
        "задание, тест-план, руководство по участию в разработке, "
        "CHANGELOG. Все необходимые для защиты артефакты собраны в "
        "ZIP-архиве, исходный код готов к публикации в открытом "
        "репозитории Git.")
    para(doc,
        "Направления развития после практики: замена локального "
        "TF-IDF-классификатора на дообученный BERT-ru для повышения "
        "качества классификации на русскоязычных письмах; интеграция с "
        "YandexGPT или GigaChat для генерации черновиков ответов; "
        "добавление веб-интерфейса поверх существующего FastAPI-бэкенда; "
        "поддержка дополнительных почтовых протоколов (POP3, Gmail API, "
        "Exchange Web Services).")
    doc.add_page_break()


def references(doc: Document) -> None:
    heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=1)
    refs = [
        "ГОСТ 34.602-2020. Информационная технология. Комплекс стандартов "
        "на автоматизированные системы. Техническое задание на создание "
        "автоматизированной системы. — М.: Стандартинформ, 2020. — 18 с.",
        "ГОСТ Р 7.0.100-2018. Библиографическая запись. Библиографическое "
        "описание. Общие требования и правила составления. — М.: "
        "Стандартинформ, 2018. — 124 с.",
        "ГОСТ 7.32-2017. Отчёт о научно-исследовательской работе. "
        "Структура и правила оформления. — М.: Стандартинформ, 2017. — 32 с.",
        "Программа преддипломной практики / ЧОУВО «МУ им. С.Ю. Витте». — "
        "М., 2024. — URL: https://www.muiv.ru/sveden/education/oop/ "
        "(дата обращения: 19.04.2026).",
        "Королькова И. А., Преображенский М. В., Зайцев С. А. Разработка "
        "и проектирование дизайна веб-интерфейсов UI/UX: учебное пособие. "
        "— М.: МУИВ, 2022. — 180 с.",
        "Рамальо Л. Python. К вершинам мастерства / пер. с англ. — "
        "2-е изд. — М.: ДМК Пресс, 2023. — 960 с.",
        "Лутц М. Изучаем Python: в 2 т. Т. 1 / пер. с англ. — 5-е изд. — "
        "М.: Вильямс, 2020. — 832 с.",
        "Педрегоса Ф. и др. Scikit-learn: машинное обучение на Python // "
        "Journal of Machine Learning Research. — 2011. — Т. 12. — "
        "С. 2825–2830.",
        "FastAPI documentation [Electronic resource]. — URL: "
        "https://fastapi.tiangolo.com (accessed: 19.04.2026).",
        "SQLAlchemy 2.0 documentation [Electronic resource]. — URL: "
        "https://docs.sqlalchemy.org/en/20/ (accessed: 19.04.2026).",
        "PyQt6 documentation [Electronic resource] / Riverbank Computing. — "
        "URL: https://www.riverbankcomputing.com/static/Docs/PyQt6/ "
        "(accessed: 19.04.2026).",
        "IEEE 829-2008. IEEE Standard for Software and System Test "
        "Documentation. — New York: IEEE, 2008. — 150 p.",
        "RFC 3501. INTERNET MESSAGE ACCESS PROTOCOL — VERSION 4rev1 / "
        "M. Crispin. — 2003. — URL: https://datatracker.ietf.org/doc/html/rfc3501 "
        "(accessed: 19.04.2026).",
        "RFC 5322. Internet Message Format / P. Resnick. — 2008. — URL: "
        "https://datatracker.ietf.org/doc/html/rfc5322 (accessed: 19.04.2026).",
        "Крейг М. Предметно-ориентированное проектирование в Python / "
        "пер. с англ. — М.: ДМК Пресс, 2022. — 340 с.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"{i}. {ref}")
        p.paragraph_format.first_line_indent = Cm(0)
    doc.add_page_break()


def appendix_a(doc: Document) -> None:
    heading(doc, "ПРИЛОЖЕНИЕ А. СКРИНШОТЫ ПОЛЬЗОВАТЕЛЬСКОГО ИНТЕРФЕЙСА", level=1)
    para(doc,
        "В приложении приведены скриншоты всех четырнадцати окон "
        "разработанного клиентского приложения. Скриншоты получены "
        "автоматически путём запуска окон в offscreen-режиме Qt "
        "(scripts/generate_assets.py) и отражают реальный внешний вид "
        "интерфейса.")
    captions = [
        ("01_login.png",              "Рисунок А.1 — Окно авторизации (LoginWindow)"),
        ("02_main.png",               "Рисунок А.2 — Главное окно (MainWindow) со списком писем"),
        ("03_email_detail.png",       "Рисунок А.3 — Детальное окно письма (EmailDetailWindow)"),
        ("04_settings.png",           "Рисунок А.4 — Окно настроек (SettingsWindow)"),
        ("05_admin.png",              "Рисунок А.5 — Панель администратора (AdminWindow)"),
        ("06_profile.png",            "Рисунок А.6 — Личный кабинет (ProfileWindow)"),
        ("07_rules.png",              "Рисунок А.7 — Окно правил (RulesWindow)"),
        ("08_reports.png",            "Рисунок А.8 — Экспорт отчётов (ReportsWindow)"),
        ("09_help.png",               "Рисунок А.9 — Встроенная справка (HelpWindow)"),
        ("10_about.png",              "Рисунок А.10 — О программе (AboutWindow)"),
        ("11_stats.png",              "Рисунок А.11 — Статистика (StatsWindow)"),
        ("12_search.png",             "Рисунок А.12 — Семантический поиск (SearchWindow)"),
        ("13_add_rule_dialog.png",    "Рисунок А.13 — Диалог добавления правила (AddRuleDialog)"),
        ("14_import_dialog.png",      "Рисунок А.14 — Диалог импорта писем (ImportDialog)"),
    ]
    for image, caption in captions:
        figure(doc, image, caption, width_cm=14)
        doc.add_paragraph()
    doc.add_page_break()


def appendix_b(doc: Document) -> None:
    heading(doc, "ПРИЛОЖЕНИЕ Б. ЛИСТИНГИ ИСХОДНОГО КОДА", level=1)
    para(doc,
        "В приложении приведены листинги всех ключевых модулей проекта "
        "без сокращений. Оригинальный исходный код также доступен в "
        "ZIP-архиве, приложенном к отчёту, и в Git-репозитории проекта.")
    para(doc,
        "Оформление листингов по ГОСТ 7.32-2017: шрифт Courier New 10 пт "
        "(моноширинный), с заголовком вида «Листинг Б.N — относительный "
        "путь до файла».")

    for i, (relpath, source_path) in enumerate(LISTING_FILES, 1):
        code_listing(doc, f"Б.{i}", relpath, source_path)


# ---------- сборка ----------


def build() -> Path:
    doc = Document()
    setup_styles(doc)
    title_page(doc)
    toc(doc)
    introduction(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    conclusion(doc)
    references(doc)
    appendix_a(doc)
    appendix_b(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


def main() -> int:
    path = build()
    size_mb = path.stat().st_size / 1024 / 1024
    print(f"Report saved: {path}")
    print(f"Size: {size_mb:.2f} MB")
    return 0


if __name__ == "__main__":
    sys.exit(main())
