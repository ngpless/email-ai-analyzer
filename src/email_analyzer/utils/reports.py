"""Экспорт результатов анализа в .docx и .xlsx.

Требование методички: «Реализовать возможность формирования документов
в форматах .docx, .xlsx и т.п. из разрабатываемого приложения.»
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from openpyxl import Workbook


def export_emails_to_docx(
    emails: Iterable[dict],
    output_path: Path,
    title: str = "Отчёт об анализе почтовых сообщений",
) -> Path:
    """Сформировать .docx со списком писем и их классификацией.

    Каждый элемент `emails` — словарь с ключами: subject, sender, sent_at,
    category, is_spam, summary.
    """
    doc = Document()
    doc.add_heading(title, level=0)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light List Accent 1"
    header_cells = table.rows[0].cells
    header_cells[0].text = "От"
    header_cells[1].text = "Тема"
    header_cells[2].text = "Дата"
    header_cells[3].text = "Категория"
    header_cells[4].text = "Спам?"

    count = 0
    for email in emails:
        cells = table.add_row().cells
        cells[0].text = str(email.get("sender", ""))
        cells[1].text = str(email.get("subject", ""))
        cells[2].text = str(email.get("sent_at", ""))
        cells[3].text = str(email.get("category", ""))
        cells[4].text = "да" if email.get("is_spam") else "нет"
        count += 1

    doc.add_paragraph()
    doc.add_paragraph(f"Всего писем в отчёте: {count}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def export_emails_to_xlsx(
    emails: Iterable[dict],
    output_path: Path,
    sheet_title: str = "Emails",
) -> Path:
    wb = Workbook()
    ws = wb.active
    if ws is None:
        raise RuntimeError("failed to create xlsx worksheet")
    ws.title = sheet_title

    headers = [
        "От", "Кому", "Тема", "Дата",
        "Категория", "Уверенность", "Спам", "Фишинг",
    ]
    ws.append(headers)

    for email in emails:
        ws.append([
            email.get("sender", ""),
            email.get("recipient", ""),
            email.get("subject", ""),
            str(email.get("sent_at", "")),
            email.get("category", ""),
            email.get("confidence", ""),
            "да" if email.get("is_spam") else "нет",
            "да" if email.get("is_phishing") else "нет",
        ])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    return output_path
